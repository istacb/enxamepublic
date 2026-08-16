#!/usr/bin/env python3
"""
MultimodalProcessor — Processador Multimodal Avançado
=====================================================
Processa:
- OCR (Tesseract, PDF)
- Análise de imagens (vision models)
- Documentos (PDF, DOCX, imagens)
- Fotos (metadados, análise visual)
- Extração de texto estruturado
"""

from __future__ import annotations

import asyncio
import base64
import hashlib
import json
import logging
import mimetypes
import os
import tempfile
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from core.ollama.client import OllamaClient, OllamaGenerateRequest

logger = logging.getLogger("enxame.multimodal")


@dataclass(slots=True)
class ProcessedFile:
    """Resultado do processamento de arquivo."""
    file_id: str
    filename: str
    mime_type: str
    size_bytes: int
    content: str = ""  # Texto extraído
    metadata: dict[str, Any] = field(default_factory=dict)
    images: list[dict[str, Any]] = field(default_factory=list)  # Para PDFs com imagens
    ocr_confidence: float = 0.0
    processing_time_ms: float = 0.0
    error: str | None = None


@dataclass(slots=True)
class ImageAnalysis:
    """Resultado de análise de imagem."""
    description: str
    objects: list[str] = field(default_factory=list)
    text_content: str = ""  # OCR da imagem
    colors: list[str] = field(default_factory=list)
    scene_type: str = ""
    confidence: float = 0.0
    metadata: dict[str, Any] = field(default_factory=dict)


class MultimodalProcessor:
    """
    Processador multimodal para OCR, imagens, documentos e fotos.
    
    Capacidades:
    - OCR via Tesseract (local) + vision models (fallback)
    - Extração de texto de PDF (PyMuPDF, pdfplumber)
    - Análise de imagens via vision models (LLaVA, BakLLaVA, etc.)
    - Processamento de DOCX, imagens, fotos
    - Metadados EXIF de fotos
    - Extração estruturada (tabelas, formulários)
    """

    def __init__(
        self,
        data_dir: Path,
        ollama_url: str = "http://localhost:11434",
        models: list[str] | None = None,
    ) -> None:
        self.data_dir = Path(data_dir)
        self.ollama_url = ollama_url
        self.models = models or []
        
        # Diretórios
        self.upload_dir = self.data_dir / "uploads"
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir = self.data_dir / "processed"
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        
        # Cliente Ollama
        self.ollama = OllamaClient(ollama_url)
        
        # Cache de processamento
        self._cache: dict[str, ProcessedFile] = {}
        
        # Verificar capacidades disponíveis
        self._has_tesseract = self._check_tesseract()
        self._has_vision_model = self._check_vision_model()
        self._has_pymupdf = self._check_pymupdf()
        self._has_pdfplumber = self._check_pdfplumber()
        self._has_pil = self._check_pil()
        
        logger.info(f"MultimodalProcessor: tesseract={self._has_tesseract}, vision={self._has_vision_model}, pymupdf={self._has_pymupdf}")

    async def initialize(self) -> None:
        """Inicializa processador."""
        # Selecionar melhor modelo de visão
        self.vision_model = await self._select_vision_model()
        logger.info(f"Modelo de visão selecionado: {self.vision_model}")

    async def close(self) -> None:
        """Fecha conexões."""
        await self.ollama.close()

    # =========================================================================
    # Verificação de Capacidades
    # =========================================================================

    def _check_tesseract(self) -> bool:
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False

    def _check_vision_model(self) -> bool:
        vision_keywords = ["llava", "bakllava", "moondream", "vision", "minicpm-v", "qwen2-vl"]
        return any(any(kw in m.lower() for kw in vision_keywords) for m in self.models)

    def _check_pymupdf(self) -> bool:
        try:
            import fitz
            return True
        except Exception:
            return False

    def _check_pdfplumber(self) -> bool:
        try:
            import pdfplumber
            return True
        except Exception:
            return False

    def _check_pil(self) -> bool:
        try:
            from PIL import Image
            return True
        except Exception:
            return False

    async def _select_vision_model(self) -> str | None:
        """Seleciona melhor modelo de visão disponível."""
        vision_models = [
            "llava:13b", "llava:7b", "bakllava:7b", 
            "moondream:1.8b", "minicpm-v:8b", "qwen2-vl:7b",
            "llava:34b", "llava-phi3:3.8b"
        ]
        available = set(self.models)
        for model in vision_models:
            if model in available:
                return model
        # Fallback: qualquer modelo que possa ter visão
        for model in self.models:
            if any(kw in model.lower() for kw in ["llava", "vision", "vl"]):
                return model
        return None

    # =========================================================================
    # API Principal
    # =========================================================================

    async def process_attachments(self, attachments: list[dict], profile: Any) -> dict[str, Any]:
        """
        Processa anexos de uma mensagem.
        
        Args:
            attachments: Lista de {filename, content (base64), content_type}
            profile: Perfil do agente que processará
        
        Returns:
            Dict com contexto multimodal para o agente
        """
        results = {
            "files": [],
            "images": [],
            "documents": [],
            "ocr_text": "",
            "image_analyses": [],
            "metadata": {},
        }
        
        for att in attachments:
            try:
                file_result = await self._process_attachment(att)
                results["files"].append(file_result)
                
                if file_result.mime_type.startswith("image/"):
                    results["images"].append(file_result)
                    if file_result.content:
                        results["ocr_text"] += f"\n[OCR {file_result.filename}]: {file_result.content}"
                elif file_result.mime_type == "application/pdf":
                    results["documents"].append(file_result)
                    if file_result.content:
                        results["ocr_text"] += f"\n[PDF {file_result.filename}]: {file_result.content}"
                else:
                    if file_result.content:
                        results["ocr_text"] += f"\n[{file_result.filename}]: {file_result.content}"
                        
            except Exception as e:
                logger.error(f"Erro processando {att.get('filename', 'unknown')}: {e}")
                results["files"].append({
                    "filename": att.get("filename", "unknown"),
                    "error": str(e),
                })
        
        return results

    async def _process_attachment(self, attachment: dict) -> ProcessedFile:
        """Processa um único anexo."""
        filename = attachment.get("filename", "unknown")
        content_b64 = attachment.get("content", "")
        content_type = attachment.get("content_type", "")
        
        # Decodificar base64
        try:
            content_bytes = base64.b64decode(content_b64)
        except Exception:
            content_bytes = content_b64.encode() if isinstance(content_b64, str) else b""
        
        # Cache key
        file_hash = hashlib.sha256(content_bytes).hexdigest()[:16]
        cache_key = f"{filename}_{file_hash}"
        
        if cache_key in self._cache:
            return self._cache[cache_key]
        
        # Salvar temporariamente
        temp_path = self.upload_dir / f"{file_hash}_{filename}"
        temp_path.write_bytes(content_bytes)
        
        try:
            # Processar por tipo MIME
            mime_type = content_type or mimetypes.guess_type(filename)[0] or "application/octet-stream"
            
            if mime_type.startswith("image/"):
                result = await self._process_image(temp_path, filename, mime_type, content_bytes)
            elif mime_type == "application/pdf":
                result = await self._process_pdf(temp_path, filename, content_bytes)
            elif mime_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",):
                result = await self._process_docx(temp_path, filename, content_bytes)
            elif mime_type.startswith("text/"):
                result = await self._process_text(temp_path, filename, mime_type, content_bytes)
            else:
                result = ProcessedFile(
                    file_id=file_hash,
                    filename=filename,
                    mime_type=mime_type,
                    size_bytes=len(content_bytes),
                    error=f"Tipo não suportado: {mime_type}",
                )
            
            # Cache
            self._cache[cache_key] = result
            return result
            
        finally:
            # Limpar temp
            try:
                temp_path.unlink()
            except Exception:
                pass

    # =========================================================================
    # Processadores por Tipo
    # =========================================================================

    async def _process_image(self, path: Path, filename: str, mime_type: str, content_bytes: bytes) -> ProcessedFile:
        """Processa imagem: OCR + análise visual."""
        start = asyncio.get_event_loop().time()
        
        # OCR via Tesseract
        ocr_text = ""
        ocr_confidence = 0.0
        if self._has_tesseract:
            try:
                import pytesseract
                from PIL import Image
                img = Image.open(path)
                ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
                ocr_text = " ".join([w for w in ocr_data["text"] if w.strip()])
                confidences = [c for c in ocr_data["conf"] if c > 0]
                ocr_confidence = sum(confidences) / len(confidences) / 100.0 if confidences else 0.0
            except Exception as e:
                logger.warning(f"OCR falhou: {e}")
        
        # Análise visual via vision model
        analysis = None
        if self.vision_model:
            analysis = await self._analyze_image_with_vision(path, content_bytes)
        
        # Metadados EXIF
        metadata = {}
        if self._has_pil:
            try:
                from PIL import Image
                img = Image.open(path)
                metadata = {
                    "width": img.width,
                    "height": img.height,
                    "mode": img.mode,
                    "format": img.format,
                }
                if hasattr(img, "_getexif") and img._getexif():
                    exif = img._getexif()
                    metadata["exif"] = {k: str(v) for k, v in exif.items() if v}
            except Exception:
                pass
        
        return ProcessedFile(
            file_id=hashlib.sha256(content_bytes).hexdigest()[:16],
            filename=filename,
            mime_type=mime_type,
            size_bytes=len(content_bytes),
            content=ocr_text or (analysis.text_content if analysis else ""),
            metadata={**metadata, "vision_analysis": analysis.__dict__ if analysis else None},
            ocr_confidence=ocr_confidence,
            processing_time_ms=(asyncio.get_event_loop().time() - start) * 1000,
        )

    async def _analyze_image_with_vision(self, path: Path, content_bytes: bytes) -> ImageAnalysis:
        """Analisa imagem usando modelo de visão."""
        try:
            # Codificar para base64
            b64 = base64.b64encode(content_bytes).decode()
            
            prompt = """Analise esta imagem e retorne JSON com:
{
  "description": "descrição detalhada da imagem",
  "objects": ["objeto1", "objeto2"],
  "text_content": "texto visível na imagem (OCR)",
  "colors": ["cor1", "cor2"],
  "scene_type": "tipo de cena (documento, foto, screenshot, diagrama, etc.)",
  "confidence": 0.9
}"""
            
            from core.ollama.client import OllamaGenerateRequest
            resp = await self.ollama.generate(
                OllamaGenerateRequest(
                    model=self.vision_model,
                    prompt=prompt,
                    images=[b64],
                    temperature=0.1,
                    num_ctx=4096,
                )
            )
            
            import json
            data = json.loads(resp.strip())
            return ImageAnalysis(**data)
        except Exception as e:
            logger.warning(f"Análise de visão falhou: {e}")
            return ImageAnalysis(description="", confidence=0.0)

    async def _process_pdf(self, path: Path, filename: str, content_bytes: bytes) -> ProcessedFile:
        """Processa PDF: extração de texto + imagens + OCR."""
        start = asyncio.get_event_loop().time()
        
        text_content = ""
        images = []
        metadata = {"pages": 0}
        
        # Tentar PyMuPDF (fitz)
        if self._has_pymupdf:
            try:
                import fitz
                doc = fitz.open(path)
                metadata["pages"] = len(doc)
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    # Texto
                    text_content += page.get_text()
                    
                    # Imagens
                    img_list = page.get_images(full=True)
                    for img_index, img in enumerate(img_list):
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n < 5:  # GRAY or RGB
                            img_bytes = pix.tobytes("png")
                            img_b64 = base64.b64encode(img_bytes).decode()
                            images.append({
                                "page": page_num + 1,
                                "index": img_index,
                                "data": img_b64,
                                "width": pix.width,
                                "height": pix.height,
                            })
                        pix = None
                doc.close()
            except Exception as e:
                logger.warning(f"PyMuPDF falhou: {e}")
        
        # Fallback: pdfplumber
        if not text_content and self._has_pdfplumber:
            try:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    metadata["pages"] = len(pdf.pages)
                    for page in pdf.pages:
                        text_content += page.extract_text() or ""
            except Exception as e:
                logger.warning(f"pdfplumber falhou: {e}")
        
        # OCR nas imagens extraídas se não há texto
        ocr_confidence = 0.0
        if not text_content.strip() and images and self._has_tesseract:
            try:
                import pytesseract
                from PIL import Image
                import io
                ocr_texts = []
                confidences = []
                for img_data in images[:5]:  # Limitar a 5 imagens
                    img_bytes = base64.b64decode(img_data["data"])
                    img = Image.open(io.BytesIO(img_bytes))
                    ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
                    page_text = " ".join([w for w in ocr_data["text"] if w.strip()])
                    if page_text:
                        ocr_texts.append(page_text)
                    confs = [c for c in ocr_data["conf"] if c > 0]
                    if confs:
                        confidences.extend(confs)
                if ocr_texts:
                    text_content = "\n".join(ocr_texts)
                if confidences:
                    ocr_confidence = sum(confidences) / len(confidences) / 100.0
            except Exception as e:
                logger.warning(f"OCR de PDF falhou: {e}")
        
        return ProcessedFile(
            file_id=hashlib.sha256(content_bytes).hexdigest()[:16],
            filename=filename,
            mime_type="application/pdf",
            size_bytes=len(content_bytes),
            content=text_content,
            metadata=metadata,
            images=images,
            ocr_confidence=ocr_confidence,
            processing_time_ms=(asyncio.get_event_loop().time() - start) * 1000,
        )

    async def _process_docx(self, path: Path, filename: str, content_bytes: bytes) -> ProcessedFile:
        """Processa DOCX."""
        start = asyncio.get_event_loop().time()
        
        try:
            from docx import Document
            doc = Document(path)
            text_content = "\n".join([p.text for p in doc.paragraphs])
            
            # Tabelas
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "tables_data": tables,
            }
            
            return ProcessedFile(
                file_id=hashlib.sha256(content_bytes).hexdigest()[:16],
                filename=filename,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                size_bytes=len(content_bytes),
                content=text_content,
                metadata=metadata,
                processing_time_ms=(asyncio.get_event_loop().time() - start) * 1000,
            )
        except Exception as e:
            logger.error(f"DOCX processing falhou: {e}")
            return ProcessedFile(
                file_id=hashlib.sha256(content_bytes).hexdigest()[:16],
                filename=filename,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                size_bytes=len(content_bytes),
                error=str(e),
            )

    async def _process_text(self, path: Path, filename: str, mime_type: str, content_bytes: bytes) -> ProcessedFile:
        """Processa arquivo de texto."""
        try:
            text = content_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = content_bytes.decode("latin-1")
            except Exception:
                text = content_bytes.decode("utf-8", errors="replace")
        
        return ProcessedFile(
            file_id=hashlib.sha256(content_bytes).hexdigest()[:16],
            filename=filename,
            mime_type=mime_type,
            size_bytes=len(content_bytes),
            content=text,
            processing_time_ms=1.0,
        )

    async def process_files(self, files: list[dict]) -> dict[str, Any]:
        """Processa múltiplos arquivos (para endpoint de upload)."""
        results = []
        for f in files:
            result = await self._process_attachment(f)
            results.append({
                "file_id": result.file_id,
                "filename": result.filename,
                "mime_type": result.mime_type,
                "size_bytes": result.size_bytes,
                "content": result.content[:5000] if result.content else "",  # Limitar
                "metadata": result.metadata,
                "ocr_confidence": result.ocr_confidence,
                "error": result.error,
            })
        return {"files": results, "total": len(results)}

    def get_stats(self) -> dict[str, Any]:
        """Estatísticas do processador."""
        return {
            "capabilities": {
                "tesseract_ocr": self._has_tesseract,
                "vision_model": self._has_vision_model,
                "vision_model_name": self.vision_model,
                "pymupdf": self._has_pymupdf,
                "pdfplumber": self._has_pdfplumber,
                "pil": self._has_pil,
            },
            "supported_types": [
                "image/*", "application/pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "text/*",
            ],
            "cache_size": len(self._cache),
        }