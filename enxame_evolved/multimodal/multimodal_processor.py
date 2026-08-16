"""
MultiModalProcessor — Processamento Multimodal
==============================================
OCR, análise de imagens, documentos, fotos.
Integra com bibliotecário existente e Ollama vision models.
"""

from __future__ import annotations

import base64
import io
import json
import logging
import mimetypes
import os
import tempfile
import uuid
from contextlib import suppress
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import httpx

# OCR
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

# PDF
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except Exception:
    PDF_AVAILABLE = False

# Office docs
try:
    import docx
    import openpyxl
    OFFICE_AVAILABLE = True
except Exception:
    OFFICE_AVAILABLE = False

logger = logging.getLogger("enxame.multimodal")


@dataclass(slots=True)
class ProcessedAttachment:
    """Resultado de processamento de anexo."""
    attachment_id: str
    original_filename: str
    mime_type: str
    content_type: str  # 'text', 'image', 'document', 'photo'
    extracted_text: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)
    analysis: str = ""
    error: str | None = None


@dataclass(slots=True)
class MultiModalResult:
    """Resultado consolidado de processamento multimodal."""
    attachments: list[ProcessedAttachment]
    combined_text: str = ""
    summary: str = ""
    errors: list[str] = field(default_factory=list)


class MultiModalProcessor:
    """
    Processador multimodal para OCR, imagens, documentos e fotos.
    
    Capacidades:
    - OCR (Tesseract) para imagens e PDFs escaneados
    - Análise de imagens via Ollama vision models (llava, bakllava, etc.)
    - Extração de texto de PDFs (PyMuPDF)
    - Processamento de documentos Office (docx, xlsx)
    - Análise de fotos (metadados, descrição)
    """
    
    SUPPORTED_IMAGE_TYPES = {
        'image/jpeg', 'image/png', 'image/gif', 'image/webp', 
        'image/bmp', 'image/tiff', 'image/heic'
    }
    SUPPORTED_DOC_TYPES = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # docx
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',  # xlsx
        'application/msword',  # doc
        'application/vnd.ms-excel',  # xls
        'text/plain',
        'text/markdown',
        'text/csv',
    }
    
    def __init__(
        self,
        ollama_url: str = "http://localhost:11434",
        model: str = "llama3.2:3b",
        vision_model: str = "llava:7b",
        cache_dir: Path | None = None,
    ) -> None:
        self.ollama_url = ollama_url.rstrip('/')
        self.model = model
        self.vision_model = vision_model
        self.cache_dir = Path(cache_dir) if cache_dir else Path(tempfile.gettempdir()) / "enxame_multimodal"
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        self._client = httpx.AsyncClient(timeout=180.0)
        self._vision_model_available = False
        
    async def initialize(self) -> None:
        """Verifica modelos disponíveis no Ollama."""
        try:
            resp = await self._client.get(f"{self.ollama_url}/api/tags")
            if resp.status_code == 200:
                models = [m["name"] for m in resp.json().get("models", [])]
                self._vision_model_available = any(
                    "llava" in m.lower() or "bakllava" in m.lower() or "vision" in m.lower() 
                    for m in models
                )
                if not self._vision_model_available:
                    logger.warning("Modelo vision não encontrado. Instale: ollama pull llava:7b")
        except Exception as e:
            logger.warning(f"Não foi possível verificar modelos Ollama: {e}")
    
    async def close(self) -> None:
        await self._client.aclose()
    
    # =========================================================================
    # API Principal
    # =========================================================================
    
    async def process_attachments(
        self, 
        attachments: list[dict[str, Any]], 
        task_context: str = ""
    ) -> str:
        """
        Processa lista de anexos e retorna texto consolidado.
        
        Args:
            attachments: Lista de dicts com {filename, content_base64, mime_type} 
                        ou {filepath, mime_type}
            task_context: Contexto da tarefa para guiar análise
            
        Returns:
            Texto consolidado com extrações e análises
        """
        processed = []
        
        for i, att in enumerate(attachments):
            try:
                result = await self._process_single_attachment(att, task_context)
                processed.append(result)
            except Exception as e:
                logger.error(f"Erro ao processar anexo {i}: {e}")
                processed.append(ProcessedAttachment(
                    attachment_id=str(uuid.uuid4()),
                    original_filename=att.get("filename", f"attachment_{i}"),
                    mime_type=att.get("mime_type", "unknown"),
                    content_type="error",
                    error=str(e),
                ))
        
        return self._consolidate_results(processed, task_context)
    
    async def _process_single_attachment(
        self, 
        attachment: dict[str, Any], 
        task_context: str
    ) -> ProcessedAttachment:
        """Processa um único anexo."""
        # Determinar source
        if "content_base64" in attachment:
            content = base64.b64decode(attachment["content_base64"])
            filename = attachment.get("filename", "unknown")
            mime_type = attachment.get("mime_type", self._guess_mime(filename))
        elif "filepath" in attachment:
            filepath = Path(attachment["filepath"])
            content = filepath.read_bytes()
            filename = filepath.name
            mime_type = attachment.get("mime_type", self._guess_mime(filename))
        else:
            raise ValueError("Anexo deve ter 'content_base64' ou 'filepath'")
        
        attachment_id = str(uuid.uuid4())
        
        # Classificar tipo de conteúdo
        content_type = self._classify_content_type(mime_type)
        
        # Processar baseado no tipo
        if content_type == "image":
            return await self._process_image(attachment_id, filename, mime_type, content, task_context)
        elif content_type == "document":
            return await self._process_document(attachment_id, filename, mime_type, content, task_context)
        elif content_type == "photo":
            return await self._process_photo(attachment_id, filename, mime_type, content, task_context)
        else:
            return ProcessedAttachment(
                attachment_id=attachment_id,
                original_filename=filename,
                mime_type=mime_type,
                content_type="unknown",
                error=f"Tipo não suportado: {mime_type}",
            )
    
    def _classify_content_type(self, mime_type: str) -> str:
        """Classifica tipo de conteúdo pelo MIME."""
        if mime_type in self.SUPPORTED_IMAGE_TYPES:
            return "image"
        elif mime_type in self.SUPPORTED_DOC_TYPES:
            if mime_type == "application/pdf":
                return "document"
            return "document"
        return "unknown"
    
    def _guess_mime(self, filename: str) -> str:
        mime, _ = mimetypes.guess_type(filename)
        return mime or "application/octet-stream"
    
    # =========================================================================
    # Processamento de Imagens (OCR + Vision)
    # =========================================================================
    
    async def _process_image(
        self, 
        attachment_id: str, 
        filename: str, 
        mime_type: str, 
        content: bytes,
        task_context: str
    ) -> ProcessedAttachment:
        """Processa imagem: OCR + análise vision."""
        extracted_text = ""
        analysis = ""
        metadata = {"size_bytes": len(content), "mime_type": mime_type}
        
        # 1. OCR se disponível
        if OCR_AVAILABLE:
            try:
                extracted_text = await self._ocr_image(content)
                metadata["ocr_performed"] = True
            except Exception as e:
                logger.warning(f"OCR falhou: {e}")
                metadata["ocr_performed"] = False
        else:
            metadata["ocr_performed"] = False
            metadata["ocr_reason"] = "pytesseract não instalado"
        
        # 2. Análise vision via Ollama se modelo disponível
        if self._vision_model_available:
            try:
                analysis = await self._analyze_image_vision(content, task_context)
                metadata["vision_performed"] = True
            except Exception as e:
                logger.warning(f"Vision analysis falhou: {e}")
                metadata["vision_performed"] = False
        else:
            metadata["vision_performed"] = False
            metadata["vision_reason"] = "modelo vision não disponível"
        
        return ProcessedAttachment(
            attachment_id=attachment_id,
            original_filename=filename,
            mime_type=mime_type,
            content_type="image",
            extracted_text=extracted_text,
            metadata=metadata,
            analysis=analysis,
        )
    
    async def _ocr_image(self, content: bytes) -> str:
        """OCR usando Tesseract."""
        image = Image.open(io.BytesIO(content))
        # Pré-processamento para melhor OCR
        if image.mode != 'RGB':
            image = image.convert('RGB')
        text = pytesseract.image_to_string(image, lang='por+eng')
        return text.strip()
    
    async def _analyze_image_vision(self, content: bytes, task_context: str) -> str:
        """Análise de imagem via modelo vision do Ollama."""
        b64 = base64.b64encode(content).decode()
        
        prompt = (
            f"Analise esta imagem no contexto: {task_context}\n\n"
            "Forneça uma descrição detalhada e estruturada incluindo:\n"
            "1. O que está na imagem (objetos, texto visível, cenas)\n"
            "2. Informações técnicas relevantes\n"
            "3. Qualquer texto legível (complementar ao OCR)\n"
            "4. Contexto para a tarefa em questão\n\n"
            "Responda em português brasileiro, objetivo e estruturado."
        )
        
        payload = {
            "model": self.vision_model,
            "prompt": prompt,
            "images": [b64],
            "stream": False,
            "options": {"temperature": 0.2, "num_ctx": 4096},
        }
        
        resp = await self._client.post(f"{self.ollama_url}/api/generate", json=payload)
        resp.raise_for_status()
        data = resp.json()
        return data.get("response", "").strip()
    
    # =========================================================================
    # Processamento de Documentos
    # =========================================================================
    
    async def _process_document(
        self, 
        attachment_id: str, 
        filename: str, 
        mime_type: str, 
        content: bytes,
        task_context: str
    ) -> ProcessedAttachment:
        """Processa documentos: PDF, DOCX, XLSX, TXT, MD, CSV."""
        extracted_text = ""
        metadata = {"size_bytes": len(content), "mime_type": mime_type}
        
        try:
            if mime_type == "application/pdf":
                extracted_text = await self._extract_pdf_text(content)
                metadata["extraction_method"] = "pymupdf"
            elif mime_type in {
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            }:
                extracted_text = await self._extract_docx_text(content)
                metadata["extraction_method"] = "python-docx"
            elif mime_type in {
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
            }:
                extracted_text = await self._extract_xlsx_text(content)
                metadata["extraction_method"] = "openpyxl"
            elif mime_type.startswith("text/"):
                extracted_text = content.decode("utf-8", errors="ignore")
                metadata["extraction_method"] = "direct"
            else:
                extracted_text = ""
                metadata["extraction_method"] = "unsupported"
        except Exception as e:
            logger.error(f"Erro ao extrair documento: {e}")
            metadata["extraction_error"] = str(e)
        
        return ProcessedAttachment(
            attachment_id=attachment_id,
            original_filename=filename,
            mime_type=mime_type,
            content_type="document",
            extracted_text=extracted_text[:50000],  # Limitar tamanho
            metadata=metadata,
        )
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extrai texto de PDF usando PyMuPDF."""
        if not PDF_AVAILABLE:
            return "[PDF processing não disponível - instale pymupdf]"
        
        doc = fitz.open(stream=content, filetype="pdf")
        texts = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                texts.append(f"--- Página {page_num + 1} ---\n{text}")
        doc.close()
        return "\n\n".join(texts)
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """Extrai texto de DOCX."""
        if not OFFICE_AVAILABLE:
            return "[DOCX processing não disponível - instale python-docx]"
        
        doc = docx.Document(io.BytesIO(content))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Tabelas
        for table in doc.tables:
            for row in table.rows:
                texts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(texts)
    
    async def _extract_xlsx_text(self, content: bytes) -> str:
        """Extrai texto de XLSX."""
        if not OFFICE_AVAILABLE:
            return "[XLSX processing não disponível - instale openpyxl]"
        
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        texts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            texts.append(f"--- Planilha: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None)
                if row_text:
                    texts.append(row_text)
        return "\n".join(texts)
    
    # =========================================================================
    # Processamento de Fotos (Metadados + Análise)
    # =========================================================================
    
    async def _process_photo(
        self, 
        attachment_id: str, 
        filename: str, 
        mime_type: str, 
        content: bytes,
        task_context: str
    ) -> ProcessedAttachment:
        """Processa fotos: metadados EXIF + análise vision."""
        metadata = {"size_bytes": len(content), "mime_type": mime_type}
        
        # 1. Extrair EXIF se disponível
        exif_data = {}
        try:
            image = Image.open(io.BytesIO(content))
            exif = image.getexif()
            if exif:
                for tag_id, value in exif.items():
                    tag = Image.ExifTags.TAGS.get(tag_id, tag_id)
                    exif_data[tag] = str(value)
            metadata["exif"] = exif_data
            metadata["dimensions"] = {"width": image.width, "height": image.height}
            metadata["format"] = image.format
        except Exception:
            pass
        
        # 2. Análise vision (mesmo que imagem)
        analysis = ""
        if self._vision_model_available:
            try:
                analysis = await self._analyze_image_vision(content, task_context)
                metadata["vision_performed"] = True
            except Exception as e:
                metadata["vision_error"] = str(e)
        
        return ProcessedAttachment(
            attachment_id=attachment_id,
            original_filename=filename,
            mime_type=mime_type,
            content_type="photo",
            extracted_text="",  # Fotos não têm texto extraível direto
            metadata=metadata,
            analysis=analysis,
        )
    
    # =========================================================================
    # Consolidação de Resultados
    # =========================================================================
    
    def _consolidate_results(
        self, 
        processed: list[ProcessedAttachment], 
        task_context: str
    ) -> str:
        """Consolida resultados em texto único para o LLM."""
        parts = []
        
        for att in processed:
            if att.error:
                parts.append(f"[ERRO - {att.original_filename}]: {att.error}")
                continue
            
            header = f"=== {att.original_filename} ({att.content_type.upper()}) ==="
            
            if att.extracted_text:
                # Limitar texto extraído
                text = att.extracted_text[:10000]
                parts.append(f"{header}\n[TEXTO EXTRAÍDO]:\n{text}")
            
            if att.analysis:
                parts.append(f"{header}\n[ANÁLISE VISION]:\n{att.analysis}")
            
            if att.metadata:
                meta_str = json.dumps(att.metadata, ensure_ascii=False, indent=2)[:2000]
                parts.append(f"{header}\n[METADADOS]:\n{meta_str}")
        
        if not parts:
            return "Nenhum conteúdo processável encontrado nos anexos."
        
        combined = "\n\n".join(parts)
        
        # Adicionar resumo se muitos anexos
        if len(processed) > 3:
            summary = self._generate_summary(processed)
            combined = f"[RESUMO GERAL]\n{summary}\n\n---\n\n{combined}"
        
        return combined
    
    def _generate_summary(self, processed: list[ProcessedAttachment]) -> str:
        """Gera resumo dos anexos processados."""
        by_type = {}
        for att in processed:
            if att.error:
                continue
            by_type[att.content_type] = by_type.get(att.content_type, 0) + 1
        
        parts = [f"Total de anexos processados: {len([a for a in processed if not a.error])}"]
        for ctype, count in by_type.items():
            parts.append(f"- {ctype.capitalize()}: {count}")
        
        # Verificar se há texto extraível
        has_text = any(a.extracted_text for a in processed if not a.error)
        if has_text:
            parts.append("- Texto extraível: SIM")
        
        has_vision = any(a.analysis for a in processed if not a.error)
        if has_vision:
            parts.append("- Análise visual: SIM")
        
        return "\n".join(parts)


# =============================================================================
# Função auxiliar para criar anexo a partir de arquivo local
# =============================================================================

def create_attachment_from_file(filepath: str | Path) -> dict[str, Any]:
    """Cria dict de anexo a partir de arquivo local."""
    path = Path(filepath)
    content = path.read_bytes()
    return {
        "filename": path.name,
        "content_base64": base64.b64encode(content).decode(),
        "mime_type": mimetypes.guess_type(path.name)[0] or "application/octet-stream",
    }


def create_attachment_from_bytes(
    content: bytes, 
    filename: str, 
    mime_type: str | None = None
) -> dict[str, Any]:
    """Cria dict de anexo a partir de bytes."""
    return {
        "filename": filename,
        "content_base64": base64.b64encode(content).decode(),
        "mime_type": mime_type or mimetypes.guess_type(filename)[0] or "application/octet-stream",
    }